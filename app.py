# Versionsbeschreibung V10: 
# einfache Kompetenzanalyse Letztversion KSM 
# + Überarbeitung Einschätzung Sicherheit mittels Bewertungskriterien 
# + Kurzbereicht & Detailergebnisse 
# + Niveau "nicht relevant" 
# + zusammengefasste NQR-Deskriptoren als zusätliche Entscheidungshilfe in Prompt 
# + farbliche unterlegung sicherheit (neuer parser pattern_detail)
# docx als zusätzliches Dateiupload-Format


import streamlit as st
import openai
from dotenv import load_dotenv
import os
import json
import re
import pandas as pd
import plotly.express as px
from docx import Document
from io import BytesIO
from typing import Union

load_dotenv()

client = openai.OpenAI()

st.title("Kompetenzanalyse Prototyp")

# Absoluter Pfad zur Datei
base_path = os.path.dirname(__file__)
kompetenzmodell_path = os.path.join(base_path, "kompetenzmodell_06_2025.json")

with open(kompetenzmodell_path, "r", encoding="UTF-8") as f:
    kompetenzmodell = json.load(f)


kompetenzmodell_text = json.dumps(kompetenzmodell, ensure_ascii=False, indent=2, sort_keys=True)

# kompetenzbereiche definieren
bereich_map = {
    "Kollaborationskompetenz": "Soziale Kompetenzen",
    "Kommunikationskompetenz": "Soziale Kompetenzen",
    "Resilienz": "Personale Kompetenzen",
    "Problemlösungskompetenz": "Methodenkompetenzen",
    "Lernkompetenz": "Personale Kompetenzen",
    "Serviceorientierung": "Soziale Kompetenzen",
    "Prozess- und Systemkompetenz": "Methodenkompetenzen",
    "Intrapreneurship": "Personale Kompetenzen",
    "Selbstmanagement": "Personale Kompetenzen",
    "Daten- und KI-Kompetenz": "Methodenkompetenzen",
    "Führen": "Soziale Kompetenzen",
    "Kritisches Denken": "Personale Kompetenzen"
    
}

from docx import Document
import re
from typing import Union

def parse_apbeschreibung_docx(file: Union[str, BytesIO], filename: str = ""):
    doc = Document(file)
    lines = []

    # Absätze
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())

    # Tabelleninhalte ergänzen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())

    def extract_section(start_markers, end_markers=None):
        start_index = None
        end_index = None
        for i, line in enumerate(lines):
            if any(start.lower() in line.lower() for start in start_markers) and start_index is None:
                start_index = i
        if start_index is not None and end_markers:
            for i in range(start_index + 1, len(lines)):
                if any(end.lower() in lines[i].lower() for end in end_markers):
                    end_index = i
                    break
        if start_index is not None:
            return "\n".join(lines[start_index + 1:end_index]).strip() if end_index else "\n".join(lines[start_index + 1:]).strip()
        return ""

    parsed = {
        "verwendung": "",
        "bezeichnung": extract_section(["2.", "FUNKTION DES ARBEITSPLATZES"], ["5."]),
        "aufgaben": extract_section(["5.", "AUFGABEN DES ARBEITSPLATZES"], ["6."]),
        "ziele": extract_section(["6.", "ZIELE DES ARBEITSPLATZES"], ["7."]),
        "taetigkeiten": extract_section(["7.", "KATALOG", "TÄTIGKEITEN"], ["12.", "13.", "SONSTIGE", "BESONDERE"])
    }

    if not parsed["verwendung"] and filename:
        match = re.search(r"([A-Z]\d[_/-]\d)", filename.upper())
        if match:
            parsed["verwendung"] = match.group(1).replace("-", "/")

    return parsed


uploaded_file = st.file_uploader("Arbeitsplatzbeschreibung hochladen", type=["txt", "docx"])

if uploaded_file:
    file_extension = uploaded_file.name.split('.')[-1].lower()

    if file_extension == "txt":
        text = uploaded_file.read().decode("utf-8")
    elif file_extension == "docx":
        parsed = parse_apbeschreibung_docx(uploaded_file, uploaded_file.name)
        text = f"""
Verwendung: {parsed['verwendung']}
Bezeichnung: {parsed['bezeichnung']}

Aufgaben:
{parsed['aufgaben']}

Ziele:
{parsed['ziele']}

Tätigkeiten:
{parsed['taetigkeiten']}
""".strip()
    else:
        st.error("Nur TXT und DOCX werden aktuell unterstützt.")
        st.stop()

    st.subheader("Vorprüfung: Extrahierte Inhalte aus der Arbeitsplatzbeschreibung")

    st.markdown(f"""
- **Verwendung:** {parsed['verwendung']}
- **Bezeichnung (Funktion):** {parsed['bezeichnung']}
- **Aufgaben:**  
{parsed['aufgaben']}
- **Ziele:**  
{parsed['ziele']}
- **Tätigkeiten:**  
{parsed['taetigkeiten']}
""")


    
    if st.button("Analyse starten"):
        prompt = f"""
Hier ist das Kompetenzmodell:

{kompetenzmodell_text}

Analysiere die folgende Arbeitsplatzbeschreibung:

{text}

Bestimme für alle Kompetenzen des Kompetenzmodells jeweils das niedrigste Kompetenzniveau, das **gerade noch ausreicht**, um die Tätigkeiten in der Arbeitsplatzbeschreibung **ohne wesentliche Einschränkungen** erfolgreich auszuführen.

Wenn ein niedrigeres Niveau offensichtlich **nicht ausreicht**, dann wähle das **nächsthöhere**. Wähle **kein höheres Niveau**, nur weil es **hilfreich oder wünschenswert** wäre – entscheidend ist, was **zur erfolgreichen Bewältigung notwendig** ist. Ist keines der Niveaus erforderlich und damit die gesamte Kompetenz mit Blick auf die Tätigkeiten nicht relevant gib bei Niveau "0" als Zahl an. 

Beziehe **alle angegebenen Tätigkeitsbereiche und alle konkreten Tätigkeiten** vollständig in deine Bewertung ein.

Auch wenn **nur eine** der beschriebenen Tätigkeiten ein höheres Kompetenzniveau erfordert, ist **dieses höhere Niveau ausschlaggebend** für die gesamte Bewertung dieser Kompetenz.

Die Bewertung darf sich **nicht am Durchschnitt** oder an der Anzahl der einfacheren Tätigkeiten orientieren. 
Es zählt ausschließlich, welches Niveau **mindestens notwendig** ist, um **alle** beschriebenen Tätigkeiten angemessen ausführen zu können.

Nutze folgende Orientierung zur Einschätzung des passenden Kompetenzniveaus:

**Kompetenzniveau 1 (NQR 1–3)**  
- Grundlegendes Faktenwissen, einfache praktische Tätigkeiten nach Anweisung, Arbeiten im vorstrukturierten Kontext.  
- Begrenzte Eigenverantwortung, geringe Selbstständigkeit.

**Kompetenzniveau 2 (NQR 4–5)**  
- Breites Theorie-Faktenwissen, eigenständige Problemlösung mit bekannten Methoden, Verantwortung für Arbeitsaufgaben.  
- Routinetätigkeiten und Verbesserung von Abläufen, Selbstständigkeit innerhalb klarer Grenzen.

**Kompetenzniveau 3 (NQR 6–8)**  
- Spezialisiertes bzw. forschungsbasiertes Wissen, kreative oder systemische Problemlösung, hohe Eigenverantwortung.  
- Innovationsfähigkeit, Führung und Verantwortung für komplexe Aufgaben und andere Personen.


Begründe deine Entscheidung klar: 
Welche Anforderung(en) in der Arbeitsplatzbeschreibung führen zur Wahl dieses Niveaus.

Gib bei jeder Kompetenz für die Einschätzung des Niveaus eine Antwortsicherheit zwischen 1 (sehr unsicher) und 5 (sehr sicher) an. Begründe diese Einschätzung. Beziehe dich dabei auf folgende Kriterien:
- Wie klar ist die Tätigkeitsbeschreibung im Hinblick auf diese Kompetenz?
- Wie gut passt die Beschreibung zu den Merkmalen der Kompetenzniveaus?
- Gibt es alternative Interpretationen, die zu einem anderen Niveau führen könnten? 

Formatiere die Ausgabe so (keine andere Ausgabe wie z.B. zusammenfassende Tabellen ganz am Ende deines Outputs): 
**Kompetenz:** [Name] 
**Niveau:** [Zahl]
**Begründung:** [Text]
**Antwortsicherheit:** [Zahl]
**Begründung Antwortsicherheit:** [Text]
"""


        response = client.chat.completions.create(
            model="gpt-4.1-2025-04-14",
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=3000
        )

        gpt_output = response.choices[0].message.content


        # Diagrammdaten aus GPT-Output extrahieren

        pattern = r"\*\*Kompetenz:\*\*\s*(.*?)\n\*\*Niveau:\*\*\s*(\d+)"

        matches = re.findall(pattern, gpt_output)

        if matches:
            kompetenzen = []
            niveaus = []
            bereiche = []

            for match in matches: 
                kompetenz = match[0].strip()
                niveau = int(match[1])
                bereich = bereich_map.get(kompetenz, "Sonstige Kompetenzen") # bereich aus mapping ermitteln

                kompetenzen.append(kompetenz)
                niveaus.append(niveau)
                bereiche.append(bereich)
            data = pd.DataFrame({
                "Kompetenz": kompetenzen,
                "Ausprägung": niveaus,
                "Bereich": bereiche
            })

            st.header("Vorgeschlagenes Kompetenzprofil für diesen Arbeitsplatz")

            st.subheader("Kurzübersicht - Kompetenzen und vorgeschlagene Ausprägungen")

            fig = px.bar_polar(data, r="Ausprägung", theta="Kompetenz", color="Bereich",
                                color_discrete_map={
                                    "Methodenkompetenzen": "red",
                                    "Soziale Kompetenzen": "green",
                                    "Personale Kompetenzen": "blue",
                                    "Sonstige Kompetenzen": "grey"
                                })
            
            #fig.update_traces(width=30)
            fig.update_layout(polar=dict(
                radialaxis=dict(showticklabels=False, showgrid=False, showline=False, range=[0,3]),
                angularaxis=dict(showticklabels=True, showgrid=False)
                )
            )

            st.plotly_chart(fig)

            for bereich in data["Bereich"].unique():
                st.markdown(f"**{bereich}**")
                for _, row in data[data["Bereich"] == bereich].iterrows():
                    st.markdown(f"{row['Kompetenz']}: {row['Ausprägung']}")

            st.header("Detailergebnisse der Kompetenzanalyse:")

            
            pattern_detail = r"\*\*Kompetenz:\*\*\s*(.*?)\s+\*\*Niveau:\*\*\s*(\d+)\s+\*\*Begründung:\*\*\s*(.*?)\s+\*\*Antwortsicherheit:\*\*\s*(\d)\s+\*\*Begründung Antwortsicherheit:\*\*\s*(.*?)(?=\s+\*\*Kompetenz:|\Z)"


            def sicherheit_farbbox(sicherheit):
                farbe = "lightgreen" if sicherheit == 5 else "gold" if sicherheit == 4 else "salmon"
                return f"<div style='background-color:{farbe}; padding: 0.3em; border-radius: 5px;'><b>Antwortsicherheit:</b> {sicherheit}</div>"
            
            matches_detail = re.findall(pattern_detail, gpt_output, re.DOTALL)


            if matches_detail:
                for kompetenz, niveau, begruendung, sicherheit, begruendung_sicherheit in matches_detail:
                    sicherheit = int(sicherheit)
                    st.markdown(f"""
### {kompetenz}
**Niveau:** {niveau}<br>
**Begründung:** {begruendung.strip()}
{sicherheit_farbbox(sicherheit)}
<small><i>{begruendung_sicherheit.strip().rstrip("-")}</i></small><br>
""", unsafe_allow_html=True)

